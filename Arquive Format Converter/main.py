import os
import sys
from pathlib import Path

# Bibliotecas para conversões específicas
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
import pandas as pd
import openpyxl
from docx2pdf import convert

class FileConverter:
    """
    Classe para conversão entre diversos formatos de arquivo.
    Suporta conversões: txt->pdf, pdf->docx, excel->word, e outras combinações comuns.
    """
    
    def __init__(self):
        self.conversion_map = {
            ('.txt', '.pdf'): self._txt_to_pdf,
            ('.pdf', '.docx'): self._pdf_to_docx,
            ('.xlsx', '.docx'): self._excel_to_word,
            ('.xls', '.docx'): self._excel_to_word,
            ('.docx', '.pdf'): self._docx_to_pdf,
            ('.csv', '.xlsx'): self._csv_to_excel,
            ('.xlsx', '.csv'): self._excel_to_csv,
        }
    
    def convert(self, input_file, output_format=None):
        """
        Converte um arquivo para o formato especificado.
        
        Args:
            input_file (str): Caminho do arquivo de entrada
            output_format (str, opcional): Formato de saída desejado (ex: '.pdf').
                                          Se não fornecido, será inferido do nome do arquivo de saída.
        
        Returns:
            str: Caminho do arquivo convertido
        """
        input_path = Path(input_file)
        
        # Verificar se o arquivo existe
        if not input_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {input_file}")
        
        input_ext = input_path.suffix.lower()
        
        # Determinar o formato de saída
        if output_format:
            output_ext = output_format.lower() if output_format.startswith('.') else f'.{output_format.lower()}'
        else:
            # Solicitar ao usuário se não for especificado
            output_ext = input("Digite a extensão de saída (ex: .pdf): ").lower()
            if not output_ext.startswith('.'):
                output_ext = f'.{output_ext}'
        
        # Verificar se a conversão é suportada
        if (input_ext, output_ext) not in self.conversion_map:
            raise ValueError(f"Conversão de {input_ext} para {output_ext} não suportada")
        
        # Criar nome do arquivo de saída
        output_file = str(input_path.with_suffix(output_ext))
        
        # Executar a conversão
        conversion_func = self.conversion_map[(input_ext, output_ext)]
        return conversion_func(input_file, output_file)
    
    def _txt_to_pdf(self, input_file, output_file):
        """Converte arquivo TXT para PDF"""
        try:
            # Criar um PDF a partir do texto
            pdf = canvas.Canvas(output_file)
            pdf.setFont("Helvetica", 12)
            
            y_position = 800  # Posição vertical inicial
            line_height = 14  # Altura da linha
            
            with open(input_file, 'r', encoding='utf-8') as file:
                for line in file:
                    if y_position < 50:  # Nova página se estiver próximo da margem inferior
                        pdf.showPage()
                        pdf.setFont("Helvetica", 12)
                        y_position = 800
                    
                    pdf.drawString(50, y_position, line.strip())
                    y_position -= line_height
            
            pdf.save()
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de TXT para PDF: {str(e)}")
    
    def _pdf_to_docx(self, input_file, output_file):
        """Converte arquivo PDF para DOCX (extração simples de texto)"""
        try:
            # Criar um documento DOCX
            doc = Document()
            
            # Extrair texto do PDF
            with open(input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    # Adicionar texto ao documento DOCX (por parágrafo)
                    for paragraph in text.split('\n'):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
            
            doc.save(output_file)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de PDF para DOCX: {str(e)}")
    
    def _excel_to_word(self, input_file, output_file):
        """Converte arquivo Excel para Word (como tabela)"""
        try:
            # Ler dados do Excel
            df = pd.read_excel(input_file)
            
            # Criar documento Word
            doc = Document()
            doc.add_heading(f"Dados de {Path(input_file).stem}", 0)
            
            # Adicionar tabela
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Preencher cabeçalhos
            for col_num, column_name in enumerate(df.columns):
                table.cell(0, col_num).text = str(column_name)
            
            # Preencher dados
            for row_num, row_data in enumerate(df.itertuples(index=False)):
                for col_num, cell_value in enumerate(row_data):
                    table.cell(row_num + 1, col_num).text = str(cell_value)
            
            doc.save(output_file)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de Excel para Word: {str(e)}")
    
    def _docx_to_pdf(self, input_file, output_file):
        """Converte arquivo DOCX para PDF"""
        try:
            convert(input_file, output_file)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de DOCX para PDF: {str(e)}")
    
    def _csv_to_excel(self, input_file, output_file):
        """Converte arquivo CSV para Excel"""
        try:
            df = pd.read_csv(input_file)
            df.to_excel(output_file, index=False)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de CSV para Excel: {str(e)}")
    
    def _excel_to_csv(self, input_file, output_file):
        """Converte arquivo Excel para CSV"""
        try:
            df = pd.read_excel(input_file)
            df.to_csv(output_file, index=False)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Erro na conversão de Excel para CSV: {str(e)}")


def main():
    """Função principal para execução via linha de comando"""
    if len(sys.argv) < 2:
        print("Uso: python converter.py <arquivo_entrada> [formato_saida]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else None
    
    converter = FileConverter()
    try:
        output_file = converter.convert(input_file, output_format)
        print(f"Arquivo convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"Erro: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()